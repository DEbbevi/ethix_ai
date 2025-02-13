import os
import zipfile
from typing import Dict, List, Tuple, Union
import logging
from docx import Document
from docx.shared import Pt, RGBColor
import markdown
import PyPDF2
import docx2txt
from pathlib import Path
from run_prompts_for_project import EthicsProcessor, ClaudeClient, ClaudeConfig
from extract_form import load_and_create_mappings
from utils import extract_text_from_files
import re
from python_docx_replace import docx_replace, docx_blocks, docx_get_keys

# Configure logging
logging.basicConfig(
    level=logging.ERROR,
    format='%(asctime)s - %(levelname)s - %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)
logger = logging.getLogger(__name__)

def clean_title_for_filename(title: str) -> str:
    """Convert title to filename-friendly format, taking first 5 words."""
    # Remove any special characters and split into words
    words = ''.join(c if c.isalnum() or c.isspace() else ' ' for c in title).split()
    # Take first 5 words and join with underscore
    return '_'.join(words[:5])

# def docx_replace(old_file: str, new_file: str, rep: Dict[str, str]) -> None:
#     """Replace text in docx file using zipfile method."""
#     logger.debug(f"Attempting to replace in {old_file} with: {rep}")
#     try:
#         with zipfile.ZipFile(old_file, 'r') as zin:
#             with zipfile.ZipFile(new_file, 'w') as zout:
#                 for item in zin.infolist():
#                     buffer = zin.read(item.filename)
#                     if item.filename == 'word/document.xml':
#                         content = buffer.decode('utf-8')
#                         logger.debug(f"Original content: {content}")
#                         for tag, replacement_text in rep.items():
#                             tag_placeholder = f"<{tag}>"
#                             logger.debug(f"Replacing {tag_placeholder} with {replacement_text}")
#                             content = content.replace(tag_placeholder, replacement_text)
#                         logger.debug(f"Modified content: {content}")
#                         buffer = content.encode('utf-8')
#                     zout.writestr(item, buffer)
#         logger.info(f"Successfully created {new_file} with replacements")
#     except Exception as e:
#         logger.error(f"Error during docx replacement: {e}")
#         raise

def docx_replace_regex(doc_obj, regex, replace):
    """Replace text in docx object using regex pattern."""
    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, regex, replace)

class DocumentGenerator:
    def __init__(self):
        self.field_mapping = load_and_create_mappings()
        
    def _get_heading_level(self, question_id: str) -> int:
        """Determine heading level based on question ID depth."""
        return len(question_id.split('.'))
    
    def _sort_responses(self, responses: Dict[str, str]) -> List[Tuple[str, str, str]]:
        """Sort responses by question ID and get titles."""
        # Define excluded keys
        excluded_keys = {
            'forskningsomrade',
            'BackgroundAndPurpose',
            'ParticipantRequirements',
            'RisksAndConsequences',
            'DataManagement',
            'SampleHandling',
            'ResultsAccess',
            'InsuranceAndCompensation'
        }
        
        sorted_items = []
        for question_id, content in responses.items():
            # Skip excluded keys and None content
            if question_id in excluded_keys or content is None:
                logger.debug(f"Skipping excluded or None content for question_id: {question_id}")
                continue
            
            # Find the title and option text from field mapping
            title = None
            option_text = None
            for field_info in self.field_mapping.values():
                if field_info and field_info.get('question_id') == question_id:
                    title = field_info.get('title', f"Question {question_id}")
                    # Check if this response matches any option
                    if field_info.get('options'):
                        try:
                            option_text = next(
                                (opt['text'] for opt in field_info['options'] 
                                 if str(opt.get('value', '')) == str(content)),
                                None
                            )
                        except (TypeError, AttributeError) as e:
                            logger.warning(f"Error processing options for {question_id}: {e}")
                            option_text = None
                    break
            
            # Use a default title if none was found
            if title is None:
                title = f"Question {question_id}"
                logger.warning(f"No title found for question_id: {question_id}")
            
            # Use option text if available, otherwise use original content
            final_content = option_text if option_text else content
            sorted_items.append((question_id, title, str(final_content)))
        
        # Sort by question ID numerically, handling potential non-numeric parts
        def sort_key(item):
            try:
                return [int(n) for n in item[0].split('.')]
            except ValueError:
                logger.warning(f"Non-numeric question ID found: {item[0]}")
                return [0]  # Default sort position for non-numeric IDs
            
        return sorted(sorted_items, key=sort_key)
    
    def generate_markdown(self, responses: Dict[str, str]) -> str:
        """Generate markdown from responses."""
        markdown_content = []
        sorted_responses = self._sort_responses(responses)
        
        for question_id, title, content in sorted_responses:
            level = self._get_heading_level(question_id)
            heading = '#' * level
            markdown_content.append(f"{heading} {title}\n\n{content}\n")
        
        return '\n'.join(markdown_content)
    
    def save_markdown(self, markdown_content: str, output_file: str = 'etikansokan_utkast.txt'):
        """Save markdown content to file."""
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(markdown_content)
        logger.info(f"Markdown saved to {output_file}")
    
    def generate_docx(self, markdown_content: str, output_file: str = 'etikansokan_utkast.docx'):
        """Convert markdown to docx and save."""
        doc = Document()
        
        # Set default font
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        
        # Convert markdown to HTML first
        html = markdown.markdown(markdown_content)
        
        # Clean up HTML tags
        html = (html
            .replace('<p>', '')
            .replace('</p>', '\n')
            .replace('<li>', '• ')
            .replace('</li>', '\n')
            .replace('<ul>', '\n')
            .replace('</ul>', '\n')
            .replace('<ol>', '\n')
            .replace('</ol>', '\n')
            .replace('<br>', '\n')
            .replace('<br/>', '\n')
            .replace('&nbsp;', ' ')
        )
        
        # Split by headers
        sections = html.split('<h')
        
        for section in sections:
            if not section.strip():
                continue
                
            # Process headers
            if section[0].isdigit():
                level = int(section[0])
                end_header = section.find('>')
                header_text = section[end_header+1:section.find('</h')]
                content = section[section.find('</h')+5:]
                
                # Add header with appropriate style
                doc.add_heading(header_text.strip(), level=level)
                
                # Add content
                if content.strip():
                    paragraphs = content.strip().split('\n')
                    for para in paragraphs:
                        if para.strip():
                            doc.add_paragraph(para.strip())
            else:
                # Just content without header
                if section.strip():
                    paragraphs = section.strip().split('\n')
                    for para in paragraphs:
                        if para.strip():
                            doc.add_paragraph(para.strip())
        
        # Save the document
        doc.save(output_file)
        logger.info(f"Document saved to {output_file}")

    def update_forskningspersonsinformation(self, responses: Dict[str, str]) -> None:
        """Update forskningspersonsinformation.docx with responses."""
        title = responses.get("1.1", "untitled")
        clean_title = clean_title_for_filename(title)
        output_file = f"FPI_{clean_title}.docx"
        
        try:
            # Load the template
            template_path = 'forms/forskningspersonsinformation_template.docx'
            if not os.path.exists(template_path):
                raise FileNotFoundError(f"Template file not found: {template_path}")
            
            doc = Document(template_path)
            keys = docx_get_keys(doc)
            logger.info(f"Keys in document: {keys}")   
            
            # Create replacements dictionary - using simple keys (without ${})
            replacements = {
                "backgroundandpurpose": responses.get("BackgroundAndPurpose", "Ge en kort men tydlig beskrivning av bakgrund och övergripande syfte med projektet. Informera om varför just den aktuella personen tillfrågas samt hur projektet har fått tillgång till uppgifter om personen som gör att denne tillfrågas."),
                "participantrequirements": responses.get("ParticipantRequirements", "Beskriv ur forskningspersonens perspektiv vad ett deltagande innebär. Vad krävs av forskningspersonen? Vilka metoder kommer att användas? Antal besök, intervjuer, enkäter, tester och tidsåtgång? Ska prover tas? Vilken sorts prover (vävnad) ska tas? Provmängd? Det ska tydligt framgå på vilket sätt undersökningsproceduren eventuellt skiljer sig från den rutinmässiga behandlingen."),
                "risksandconsequences": responses.get("RisksAndConsequences", "Ge saklig information om de följder och risker som deltagandet kan innebära. Undvik förskönande formuleringar och formuleringar som kan innebära otillbörlig påverkan. Kan deltagandet innebära obehag, smärta, känslomässiga effekter, integritetsintrång etc.? Beskriv eventuella biverkningar och andra effekter på kort och lång sikt. I förekommande fall ska det framgå hur de projektansvariga kommer att hantera de problem som kan uppstå. Kan deltagandet i projektet/studien avbrytas vid vissa effekter? Vilken möjlighet finns till uppföljande undersökning eller samtal etc.?"),
                "datamanagement": responses.get("DataManagement", "Förklara vilken information som kommer att samlas in, hur den kommer att hanteras och förvaras samt för hur lång tid. Varifrån kommer data hämtas, vilka källor kommer att användas? Kommer informationen gå att härleda till forskningspersonen? Hur kommer tillgången till uppgifterna att se ut? Hur skyddas uppgifterna?Ange ändamålen med behandlingen av personuppgifterna och den rättsliga grunden enligt EU:s dataskyddsförordning för behandlingen. Om uppgifterna kommer att överföras till ett land utanför EU och EES-området (s.k. tredjeland) eller till en internationell organisation ska detta särskilt anges. Det ska också anges om det finns ett beslut av EU-kommissionen om att landet eller organisationen kan säkerställa en adekvat skyddsnivå och i annat fall en hänvisning till lämpliga eller passande skyddsåtgärder och hur en kopia av dessa kan erhållas eller var dessa har gjorts tillgängliga."),
                "samplehandling": responses.get("SampleHandling", "Om prover kommer att sändas inom Sverige eller utomlands för analys ska det framgå. Det ska framgå om proverna ska sändas till ett annat EU/EES-land eller till ett tredje land. Kommer proverna bevaras hos mottagaren, återlämnas, avidentifieras eller förstöras? Hur lång tid kommer proverna förvaras/analyseras i Sverige eller utomlands och inom vilken tid kommer proverna återlämnas, avidentifieras eller förstöras?"),
                "resultsaccess": responses.get("ResultsAccess", "Informera om på vilket sätt forskningspersonen kan ta del av sina individuella data respektive resultatet av hela projektet/studien. Forskningspersonens möjlighet att inte behöva ta del av eventuella analysresultat bör framgå. Det bör också framgå hur projektet kommer att hantera eventuella oförutsedda fynd."),
                "insuranceandcompensation": responses.get("InsuranceAndCompensation", "Informera om vilket försäkringsskydd som gäller. Alla forskningspersoner ska ha ett heltäckande försäkringsskydd. Det ska framgå om forskningspersonen har rätt att få ersättning för förlorad arbetsinkomst eller utgifter som är kopplade till projektet. Det ska också framgå om ersättningen är skattepliktig eller inte."),
                "title": responses.get("1.1", "Forskningsprojektets titel, ska vara samma som i etikansökan")
            }
            
            # Replace all placeholders
            docx_replace(doc, **replacements)
            
            # Save the document
            doc.save(output_file)
            logger.info(f"Successfully created {output_file}")
            
        except Exception as e:
            logger.error(f"Failed to create {output_file}: {e}")
            raise

    def update_samtyckesblankett(self, responses: Dict[str, str]) -> None:
        """Update samtyckesblankett.docx with responses."""
        title = responses.get("1.1", "untitled")
        clean_title = clean_title_for_filename(title)
        output_file = f"Samtycke_{clean_title}.docx"
        
        try:
            # Load the template
            doc = Document('forms/samtyckesblankett_template.docx')
            
            # Define replacements - using simple key (without ${})
            replacements = {
                "title": responses.get("1.1", "Titel, ska vara samma som i etikansökan")
            }
            
            # Replace all placeholders
            docx_replace(doc, **replacements)

            # Add biobank information
            try:
                if int(responses.get("14.1", "0").strip()) == 1:
                    docx_blocks(doc, biobank=True)
                else:
                    docx_blocks(doc, biobank=False)
            except ValueError:
                logger.error(f"Failed to convert '14.1' to integer: {responses.get('14.1', '0')}")
                docx_blocks(doc, biobank=False)
            
            # Save the document
            doc.save(output_file)
            logger.info(f"Successfully created {output_file}")
        except Exception as e:
            logger.error(f"Failed to create {output_file}: {e}")
            raise

def generate_documentation(responses: Dict[str, str]) -> None:
    """
    Generate all documentation from the ethics application responses.
    
    Args:
        responses: Dictionary containing all responses from the ethics application
    """
    try:
        # Log the input responses for debugging
        logger.info(f"Received responses: {responses}")
        if responses is None:
            raise ValueError("Responses dictionary is None")
        if not isinstance(responses, dict):
            raise ValueError(f"Expected dict for responses, got {type(responses)}")

        # Generate documentation
        doc_generator = DocumentGenerator()

        title = responses.get("1.1", "untitled")
        clean_title = clean_title_for_filename(title)
        output_file = f"etikansokan_utkast_{clean_title}"
        
        # Generate and save markdown/docx
        logger.info("Generating markdown content...")
        markdown_content = doc_generator.generate_markdown(responses)
        logger.info("Saving markdown content...")
        doc_generator.save_markdown(markdown_content, f"{output_file}.txt")
        logger.info("Generating docx...")
        doc_generator.generate_docx(markdown_content, f"{output_file}.docx")
        
        # Update both template documents
        logger.info("Updating forskningspersonsinformation...")
        doc_generator.update_forskningspersonsinformation(responses)
        logger.info("Updating samtyckesblankett...")
        doc_generator.update_samtyckesblankett(responses)
        
        logger.info("Successfully generated all documentation")
    except Exception as e:
        logger.error(f"Error generating documentation: {e}")
        logger.error(f"Error type: {type(e)}")
        logger.error(f"Error details: {str(e)}")
        import traceback
        logger.error(f"Traceback: {traceback.format_exc()}")
        raise

def main():
    # Initialize Claude client
    config = ClaudeConfig(api_key=os.getenv("ANTHROPIC_API_KEY"))
    claude_client = ClaudeClient(config)
    processor = EthicsProcessor(claude_client)
    
    # Example usage with multiple files
    try:
        scientific_material = extract_text_from_files([
            "path/to/document1.pdf",
            "path/to/document2.docx",
            "path/to/document3.txt"
        ])
        logger.info(f"Extracted scientific material: {bool(scientific_material)}")
    except ValueError as e:
        logger.error(f"Failed to extract text from files: {e}")
        return
    
    try:
        # Process ethics application
        logger.info("Processing forskningsomrade...")
        base_response = processor.process_forskningsomrade(scientific_material)
        logger.info(f"Base response received: {bool(base_response)}")
        
        logger.info("Processing remaining prompts...")
        responses = processor.process_remaining_prompts(scientific_material)
        logger.info(f"Remaining responses received: {bool(responses)}")
        
        # Combine all responses
        all_responses = {
            "forskningsomrade": base_response,
            **responses
        }
        logger.info(f"Combined responses: {bool(all_responses)}")
        
        # Generate documentation using the new function
        generate_documentation(all_responses)
    except Exception as e:
        logger.error(f"Error processing ethics application: {e}")
        logger.error(f"Error type: {type(e)}")
        logger.error(f"Error details: {str(e)}")
        import traceback
        logger.error(f"Traceback: {traceback.format_exc()}")

if __name__ == "__main__":
    main() 